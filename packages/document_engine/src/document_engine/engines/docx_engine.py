import logging
from pathlib import Path

from docx import Document
from tabulate import tabulate

from ..schema import Fragment

logger = logging.getLogger(__name__)


class DocxEngine:
    """引擎：负责将 .docx 文档转换为 Fragment 列表"""

    def parse(self, file_path: str | Path) -> list[Fragment]:
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {path}")
            return []

        fragments = []
        try:
            doc = Document(path)

            # 由于 docx 的元素（段落、表格）是顺序排列的，我们遍历整个文档体
            for element in doc.element.body:
                if element.tag.endswith("p"):  # Paragraph
                    # para_idx = doc.element.body.index(element)
                    # 这种方法虽然底层，但能保证相对顺序
                    # 简单起见，我们直接遍历 doc.paragraphs 和 doc.tables 也是可以的
                    # 但是 doc.paragraphs 不包含表格，顺序会乱
                    pass

            # 简单的顺序处理策略：doc.paragraphs 和 doc.tables 按顺序组合
            # python-docx 没有一个统一的顺序迭代器，但我们可以通过 element 的位置来判断

            # 更简单的做法：顺序遍历所有 paragraphs。注意：docx 的表格实际上在 paragraphs 之间。
            # 为了实现简单的 v0.2.0，我们可以先遍历所有段落和表格。

            # 这里采用一个比较稳妥的顺序获取方式
            for block in self._iter_block_items(doc):
                if hasattr(block, "text"):  # Paragraph
                    if block.text.strip():
                        fragments.append(
                            Fragment(
                                type="text",
                                y0=0.0,
                                content=block.text,
                                meta={"source_file": path.name},
                            )
                        )
                else:  # Table
                    data = []
                    for row in block.rows:
                        data.append([cell.text.strip() for cell in row.cells])

                    if data:
                        md_table = tabulate(data, headers="firstrow", tablefmt="pipe")
                        fragments.append(
                            Fragment(
                                type="table",
                                y0=0.0,
                                content=md_table,
                                meta={"source_file": path.name},
                            )
                        )

        except Exception as e:
            logger.error(f"Error parsing Docx {path}: {e}")

        return fragments

    def _iter_block_items(self, parent):
        """
        遍历 Document 内部的块级元素（段落和表格）。
        """
        from docx.document import Document as _Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        else:
            raise TypeError("Parent must be a Document")

        for child in parent_elm.iterchildren():
            if isinstance(child, Paragraph) or child.tag.endswith("p"):
                yield Paragraph(child, parent)
            elif isinstance(child, Table) or child.tag.endswith("tbl"):
                yield Table(child, parent)
